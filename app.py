import os
import time
import fitz  # PyMuPDF for PDF handling
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import streamlit as st
import google.generativeai as genai
import re
from dotenv import load_dotenv
from backend import pdf_to_images, extract_pdf_content, process_page
load_dotenv()

# Default API Key
DEFAULT_API_KEY = os.getenv("API_KEY")

pdf_extraction_prompt = """
You will be given pages of a PDF file containing text in Arabic. Your task is to extract the content from each page and categorize it into the following sections in **JSON format**:

1. **Headers**:
   - Extract all text from the very top section of the page, typically found in the margin above the main content.
   - This section may include information such as document titles, section names, or repetitive information present across multiple pages. It may be present on the page may be not
   - The header is distinct from the main heading of the page.
   - Header is mostly seprated by the actual line using the line below the header.

2. **Heading**:
   - Extract the main heading of the page if it exists.
   - The main heading is typically a prominent line of text near the top of the body, often in larger, bold font and distinct in appearance from the header.
   -it is center-aligned
   - Do not confuse this with text present in the header, or main content section.

3. **Main Content**:
   - Extract all text from the body of the page that is not part of the header, heading, footer, or footnotes.

4. **Footnotes (Text Below the Black Line)**:
   - Carefully identify any black horizontal line present on the page.
   - If the line exists, categorize all text below it as "Footnotes".
   - The black line will typically cover about half the width of the page and is visually distinct.
   - If no black line is present, the "Footnotes" section should be empty for that page.

5. **Footers**:
   - Extract all text from the footer section of the page, typically located at the very bottom.
   - Footers often include repetitive elements such as page numbers or document-specific references and should not overlap with footnotes.

### Formatting Guidelines:
- **Maintain the original Arabic text formatting** as closely as possible.
- Use the following formatting rules:
  - **Headings**: Represent the main headings in a larger, bold font.
- Ensure all extracted text is **right-aligned** to match proper Arabic formatting.

### Output Format:
For each page, provide the extracted data in the following JSON structure:
{
  "header": "<Arabic text of the header>",
  "heading": "<Arabic text of the heading>",
  "main_content": "<Arabic text of the main content>",
  "footer": "<Arabic text of the footer>",
  "footnotes": "<Arabic text of the footnotes>"
}
"""

def find_and_replace_in_docx(doc, find_texts, replace_texts):
    """
    Replaces all occurrences of specified Arabic text in the document.
    """
    if len(find_texts) != len(replace_texts):
        raise ValueError("Find and Replace lists must have the same length.")

    for find_text, replace_text in zip(find_texts, replace_texts):
        for paragraph in doc.paragraphs:
            if find_text in paragraph.text:
                paragraph.text = paragraph.text.replace(find_text, replace_text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if find_text in cell.text:
                        cell.text = cell.text.replace(find_text, replace_text)

# Streamlit Sidebar for Navigation
st.sidebar.header("Navigation")
options = ["Process PDF", "Find and Replace"]
choice = st.sidebar.radio("Go to:", options)

# Process PDF Section
if choice == "Process PDF":
    st.title("Arabic PDF to Word Converter")
    st.write("Upload a PDF, extract Arabic content, and download the result in a Word document.")

    # Input fields
    user_api_key = st.text_input("Enter your Gemini API Key (optional):", type="password")
    pdf_file = st.file_uploader("Upload a PDF file", type=["pdf"])
    output_file_name = st.text_input("Enter output Word file name (with .docx extension):", "result.docx")
    start_page = st.number_input("Start Page (1-based index):",value=1)
    end_page = st.number_input("End Page (inclusive):",value=1)

    # Processing options
    footnotes = st.checkbox("Include Footnotes", value=False)
    headers = st.checkbox("Include Headers and Footers", value=False)
    extra_chars = st.text_area("Characters to Remove (comma-separated):", "").split(",")

    if st.button("Process PDF"):
        if not pdf_file:
            st.error("Please upload a PDF file.")
        else:
            try:
                # Step 1: Save the uploaded PDF
                pdf_path = os.path.join("temp", "uploaded_pdf.pdf")
                os.makedirs("temp", exist_ok=True)
                with open(pdf_path, "wb") as f:
                    f.write(pdf_file.read())

                # Step 2: Validate and enforce page limits
                pdf_document = fitz.open(pdf_path)
                total_pages = len(pdf_document)
                pdf_document.close()

                if end_page == 0 or end_page > total_pages:
                    end_page = total_pages

                if not user_api_key and (end_page - start_page + 1) > 10:
                    st.warning("API key not provided. Limiting processing to 10 pages.")
                    end_page = min(start_page + 9, total_pages)

                # Step 3: Convert PDF pages to images
                output_folder = "temp_images"
                pdf_to_images(pdf_path, output_folder, start_page=start_page, end_page=end_page)

                # Step 4: Initialize Word document
                doc = Document()

                # Step 5: Extract content and process pages
                st.write("Extracting content from the PDF...")
                
                    
                try:
                    page_content = extract_pdf_content(
                        pdf_extraction_prompt,
                        start_page=start_page,
                        end_page=end_page,
                        api_key=user_api_key if user_api_key else None
                    )
                    
                    # Process the extracted content into the Word document
                    i=1
                    for page_data in page_content:
                        st.write("Content extraction complete.",page_data)
                        try:
                            if extra_chars == [""]:
                                process_page(
                                page_data=page_data,
                                doc=doc,
                                page_number=i,
                                need_header_and_footer=headers,
                                
                                need_footnotes=footnotes,
                                )
                            else:
                                process_page(
                                page_data=page_data,
                                doc=doc,
                                page_number=i,
                                need_header_and_footer=headers,
                                need_footnotes=footnotes,
                                remove_characters=extra_chars
                                )
                        except Exception as e:
                            st.error(f"Error processing page {i}: {e}")
                            continue
                        i=i+1
                except Exception as e:
                    st.error(f"Error processing page {e}")

                # Step 6: Save the Word document
                output_path = os.path.join("temp", output_file_name)
                doc.save(output_path)

                # Step 7: Provide a download link
                with open(output_path, "rb") as f:
                    st.download_button("Download Word Document", f, file_name=output_file_name)

            except Exception as e:
                st.error(f"Error: {e}")
# Find and Replace Section
elif choice == "Find and Replace":
    # Inject CSS to align text inputs to the right
    st.markdown(
        """
        <style>
        .right-align input {
            text-align: right !important;
        }
        .stTextInput input {
            text-align: right !important;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

    st.title("Find and Replace in Arabic DOCX")
    st.write("Upload a DOCX file, specify text to find and replace, and download the updated document.")

    docx_file = st.file_uploader("Upload a DOCX file for Editing", type=["docx"])

    # Initialize session state for dynamic find-replace inputs
    if "find_replace_pairs" not in st.session_state:
        st.session_state.find_replace_pairs = [("", "")]

    st.subheader("Specify Text to Find and Replace (Use copy-paste for quick and better results)")

    # Dynamic inputs for find and replace pairs
    for i, (find_text, replace_text) in enumerate(st.session_state.find_replace_pairs):
        cols = st.columns(2)
        with cols[0]:
            st.session_state.find_replace_pairs[i] = (
                st.text_input(
                    f"Text to Find {i + 1} (Arabic):",
                    value=find_text,
                    key=f"find_{i}",
                    placeholder="Enter text to find",
                ),
                st.session_state.find_replace_pairs[i][1]
            )
        with cols[1]:
            st.session_state.find_replace_pairs[i] = (
                st.session_state.find_replace_pairs[i][0],
                st.text_input(
                    f"Replacement Text {i + 1} (Arabic):",
                    value=replace_text,
                    key=f"replace_{i}",
                    placeholder="Enter replacement text",
                )
            )

    # Button to add another pair of inputs
    if st.button("Add Another Find-Replace Pair"):
        st.session_state.find_replace_pairs.append(("", ""))

    output_file_name_edit = st.text_input("Enter output Word file name (without extension):", "مُتَجَدِّدة يَوْميًّا")
    output_file_name_edit +=".docx"
    if st.button("Perform Find and Replace"):
        if not docx_file:
            st.error("Please upload a DOCX file.")
        else:
            try:
                doc_path = os.path.join("temp", "uploaded_docx.docx")
                os.makedirs("temp", exist_ok=True)
                with open(doc_path, "wb") as f:
                    f.write(docx_file.read())

                doc = Document(doc_path)

                # Filter out empty find-replace pairs
                find_replace_pairs = [
                    (find_text.strip(), replace_text.strip())
                    for find_text, replace_text in st.session_state.find_replace_pairs
                    if find_text.strip()  # Include only valid "find" texts
                ]

                # Perform find and replace
                for paragraph in doc.paragraphs:
                    for find_text, replace_text in find_replace_pairs:
                        if find_text in paragraph.text:
                            paragraph.text = paragraph.text.replace(find_text, replace_text)

                # Save the updated document
                updated_path = os.path.join("temp", output_file_name_edit)
                doc.save(updated_path)

                # Provide download link
                with open(updated_path, "rb") as f:
                    st.download_button("Download Updated DOCX", f, file_name=output_file_name_edit)

            except Exception as e:
                st.error(f"Error processing the document: {e}")
